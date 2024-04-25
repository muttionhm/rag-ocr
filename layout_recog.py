import os
from PIL import Image
import cv2
import pytesseract
from pdf2image import convert_from_path
import time
import threading
import PyPDF2
from PIL import Image
import subprocess
import docx
Image.MAX_IMAGE_PIXELS = 3300000000

def pdf_process(pdf_path):
  result = []
  data_ocr = {}
  pdf_path = pdf_path
  pdf = convert_from_path(pdf_path, thread_count=12,dpi=200)
  bad_count = []
  for i in range(len(pdf)):
    # temp_text=pdf_fast.pages[i].extract_text()
    reader = PyPDF2.PdfReader(open(pdf_path,'rb'))
    temp_text = reader.pages[i].extract_text()
    if len(temp_text)<100:
      bad_count.append(i)
    else:
      data_ocr[i] = temp_text

  print(len(bad_count))

  def ocr_thread(file, id):
    # 使用Pillow打开图像
    # 使用pytesseract进行OCR
    text = pytesseract.image_to_string(file, lang='chi_sim+eng')
    # 将结果保存到data字典中，键为文件名，值为识别的文本
    data_ocr[id] = text
    return

  start = time.time()
  threads = 2
  process_num = len(bad_count)//threads
  for i in range(process_num):
    threadings = []
    for j in range(threads):
      thread = threading.Thread(target=ocr_thread, args=(pdf[bad_count[i*threads+j]], bad_count[i*threads+j]))
      threadings.append(thread)
      thread.start()
    for thread in threadings:
        thread.join()

  if len(bad_count)-process_num*threads>0:
    threadings = []
    for j in range (len(bad_count)-process_num*threads):
      thread = threading.Thread(target=ocr_thread, args=(pdf[bad_count[process_num*threads+j]], bad_count[process_num*threads+j]))
      threadings.append(thread)
      thread.start()
    for thread in threadings:
      thread.join()


  for i in range(len(pdf)):
    result.append(data_ocr[i])
  m_result = ''.join(result)
  return m_result

def convert_doc_to_docx(doc_path, docx_path):
    # 使用LibreOffice的soffice命令进行转换
    subprocess.run([
        'soffice',
        '--headless',
        '--convert-to',
        'docx',
        doc_path,
        '--outdir',
        os.path.dirname(docx_path)
    ])

def doc_process(doc_path):
  if doc_path.endswith('.doc'):
          # 转换文件
      doc_path = doc_path
      docx_path = ''
      convert_doc_to_docx(doc_path, docx_path)
      docx_file = doc_path[:-4]+'.docx'
      docx_file = docx_file.split('/')[-1]
      print(docx_file)
      doc = docx.Document(docx_file)
      text = [i.text for i in doc.paragraphs]
      return (''.join(text))
  else:
    doc = docx.Document(doc_path) 
    text = [i.text for i in doc.paragraphs]
    return (''.join(text))

